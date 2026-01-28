# -*- coding: utf-8 -*-
"""
四连问交互版：
1. 选文档
2. 输入起始序号
3. 输入文件名前缀（回车=空，用纯序号）
4. 选择输出目录（回车=源文件同目录下 split_result）
"""
import re, os
from docx import Document
from tkinter import Tk, filedialog

def pick_file(title='请选择文件'):
    Tk().withdraw()
    return filedialog.askopenfilename(title=title)

def pick_dir(title='请选择文件夹'):
    Tk().withdraw()
    return filedialog.askdirectory(title=title)

def save_doc(out_dir, base_name, idx, buff):
    new = Document()
    new.add_paragraph('\n'.join(buff))
    file_name = f'{base_name}{idx}.docx' if base_name else f'{idx}.docx'
    new.save(os.path.join(out_dir, file_name))

def main():
    # ① 选文档
    src_file = pick_file('请选择要拆分的 Word 文档')
    if not src_file:
        print('未选择文件，程序退出。')
        return

    # ② 起始序号
    while True:
        start_str = input('请输入起始序号（如 33）：').strip()
        if start_str.isdigit():
            start_idx = int(start_str)
            break
        print('请输入纯数字！')

    # ③ 文件名前缀
    base_name = input('请输入文件名前缀（直接回车=空，用纯序号）：').strip()

    # ④ 输出目录
    out_dir = input('请拖拽或手动输入输出目录（直接回车=源文件目录下split_result）：').strip()
    if not out_dir:
        out_dir = os.path.join(os.path.dirname(src_file), 'split_result')
    os.makedirs(out_dir, exist_ok=True)

    # ⑤ 拆分逻辑
    doc = Document(src_file)
    buff = []
    idx = start_idx - 1

    for p in doc.paragraphs:
        txt = p.text.strip()
        if re.match(r'^\d{1,3}[\uFF0E\.]', txt):
            if buff:
                idx += 1
                save_doc(out_dir, base_name, idx, buff)
            buff = [txt]
        else:
            buff.append(txt)

    if buff:
        idx += 1
        save_doc(out_dir, base_name, idx, buff)

    print(f'全部拆完！共 {idx - start_idx + 1} 个文件，保存在 → {out_dir}')

if __name__ == '__main__':
    main()