# -*- coding: utf-8 -*-
"""
行首序号+点即切：点后无需空格，正文再出现 31. 不误杀
"""
import re, os
from docx import Document
from tkinter import Tk, filedialog

def pick_file(title='请选择要拆分的 Word 文档'):
    Tk().withdraw()
    return filedialog.askopenfilename(title=title)

def pick_dir(title='请选择输出目录'):
    Tk().withdraw()
    return filedialog.askdirectory(title=title)

def main():
    src_file = pick_file()
    if not src_file:
        print('未选择文件，程序退出。')
        return

    out_dir = pick_dir() or os.path.join(os.path.dirname(src_file), 'split_result')
    os.makedirs(out_dir, exist_ok=True)

    base_name = input('请输入文件名前缀（回车=空，用纯序号）：').strip()

    doc  = Document(src_file)
    full = '\n'.join(p.text for p in doc.paragraphs)
    full = re.sub(r'\n+', '\n', full)

    # ★★★ 行首序号+点，点后无空格要求 ★★★
    pattern = re.compile(r'(?:^|\n)(\d{1,3})(?:[.\uff0e])(.*?)(?=\n\d{1,3}[.\uff0e]|$)', re.S)
    matches = list(pattern.finditer(full))

    for m in matches:
        idx = int(m.group(1))
        txt = f"{idx}.{m.group(2).strip()}"
        new = Document()
        new.add_paragraph(txt)
        file_name = f'{base_name}{idx}.docx' if base_name else f'{idx}.docx'
        new.save(os.path.join(out_dir, file_name))

    print(f'全部拆完！共 {len(matches)} 个文件 → {out_dir}')

if __name__ == '__main__':
    main()