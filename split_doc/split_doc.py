# -*- coding: utf-8 -*-
import re, os
from docx import Document

baseName = '人设40条'          # <<< 想改名字只改这里
src_file = 'D:\SMB\Vscodeproject\split_doc\source.docx'       # 原始长文档
out_dir  = 'D:\SMB\Vscodeproject\split_doc\split_result'
os.makedirs(out_dir, exist_ok=True)

doc  = Document(src_file)
buff = []
idx  = 0

for p in doc.paragraphs:
    txt = p.text.strip()
    # 匹配 1.  2.  …  100.  这样的序号
    if re.match(r'^\d{1,3}[\uFF0E\.]', txt):
        if buff:
            idx += 1
            new = Document()
            new.add_paragraph('\n'.join(buff))
            new.save(os.path.join(out_dir, f'{baseName}{idx}.docx'))
        buff = [txt]
    else:
        buff.append(txt)

# 最后一段
if buff:
    idx += 1
    new = Document()
    new.add_paragraph('\n'.join(buff))
    new.save(os.path.join(out_dir, f'{baseName}{idx}.docx'))

print(f'全部拆完！共 {idx} 个文件，保存在 {out_dir}')